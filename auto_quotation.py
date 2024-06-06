# Import docx NOT python-docx
from flask import Flask, request, render_template, jsonify, send_from_directory
import docx
from docx.shared import RGBColor
from docx.shared import Cm, Pt  #加入可調整的 word 單位
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH #處理字串的置中 
from docx.enum.text import WD_LINE_SPACING 
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import pandas as pd
import shutil
import os
import argparse
import json
import datetime
import message_to_json


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def week_calculate(yy,mm,dd):
    day = datetime.datetime(int(yy), int(mm), int(dd)).weekday()
    if day == 0:
        return "一"
    elif day == 1:
        return "二"
    elif day == 2:
        return "三"
    elif day == 3:    
        return "四"
    elif day == 4:
        return "五"
    elif day == 5:    
        return "六"
    elif day == 6:    
        return "日"

def table_mystyle(table,info_flag):
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if i%2 == 0 and info_flag:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT  
                else:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  
                for run in paragraph.runs:
                    font = run.font
                    font.name = '標楷體'
                    font._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') 

def process_quotation(quotation_data):
    #print(quotation_data)
    # Create an instance of a word document
    doc = docx.Document("src/default.docx")
    #file1 = open("quotation.json", "r")
    quotation_list = quotation_data
    #print(quotation_list)
    info = dict(quotation_list[0])
    
    #print("info:")
    #print(info)
    # Add a Title to the document
    paragraph = doc.paragraphs[0]
    paragraph.style.name.startswith('Heading')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #run.font.color.rgb = RGBColor(18, 255, 0)
    paragraph.runs[0].font.name = '標楷體'
    paragraph.runs[0].font.bold = True
    paragraph.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
    paragraph.runs[0].font.size = Pt(16)
    paragraph.runs[0].text = "報價單" 
    # print(paragraph.runs[0].text)

    # Creating a table object
    table = doc.add_table(rows=5, cols=4)

    total = 0

    # Adding heading in the 1st row of the table
    title = table.cell(0,0)
    title.text = "抬頭："
    title_content = table.cell(0,1)
    title_content.text = info["抬頭"]

    contact = table.cell(1,0)
    contact.text = "聯絡人："
    contact_content = table.cell(1,1)
    contact_content.text = info["聯絡人"]

    phone = table.cell(2,0)
    phone.text = "電話："
    phone_content = table.cell(2,1)
    phone_content.text = info["電話"]

    mail = table.cell(3,0)
    mail.text = "信箱："
    mail_content = table.cell(3,1)
    mail_content.text = info["信箱"]

    tax = table.cell(0,2)
    tax.text = "統編："
    tax_content = table.cell(0,3)
    tax_content.text = info["統編"]


    quotation_num = table.cell(1,2)
    quotation_num.text = "報價單號："
    quotation_num_content = table.cell(1,3)
    quotation_num_content.text = info["報價單號"]

    quotation_date = table.cell(2,2)
    quotation_date.text = "報價日期："
    quotation_date_content = table.cell(2,3)
    quotation_date_content.text = info["報價日期"]

    event_date = table.cell(3,2)
    event_date.text = "活動日期："
    event_date_content = table.cell(3,3)
    event_date_content.text = info["活動日期"]

    address = table.cell(4,0)
    address.text = "地址："
    address_content = table.cell(4,1)
    address_content.text = info["地址"]

    table.cell(4,1).merge(table.cell(4,2)).merge(table.cell(4,3))
    # Adding style to a table

    for i in range(5):
        table.cell(i,0).width = Cm(2.77)
        table.cell(i,1).width = Cm(6.46)
        table.cell(i,2).width = Cm(2.77)
        table.cell(i,3).width = Cm(6.46)
    for row in table.rows:
        row.height = Cm(0.82)

    table.style = 'Table Normal'
    table_mystyle(table,1)

    doc.add_paragraph("")

    for quotation in  quotation_list[1:]:
        table = doc.add_table(rows=len(quotation)-2, cols=5)

        for i in range(1,5):
            table.cell(0,0).merge(table.cell(0,i))
        

        product = table.cell(1,0)
        product.text = "產品/規格"
        number = table.cell(1,2)
        number.text = "數量"
        price = table.cell(1,3)
        price.text = "單價"
        subtotal = table.cell(1,4)
        subtotal.text = "小計"
    
        j = 0
        for k, v in quotation.items():
            if j == 0:
                table.cell(0,0).text = table.cell(0,0).text + v +" "
            if j == 1:
                table.cell(0,0).text = table.cell(0,0).text + v +"（"+ week_calculate(info["報價日期"].split("/")[0],v.split("/")[0],v.split("/")[1])+"）"
            if j == 2:
                table.cell(0,0).text = table.cell(0,0).text + k + "：" + v + "；" 
            if j == 3:
                table.cell(0,0).text = table.cell(0,0).text + k + "：" + v + " "
            
            table.cell(0,0).paragraphs[0].runs[0].font.bold = True

            if j > 3:
                product_content = table.cell(j-2,0)
                product_content.text = k
                number_content = table.cell(j-2,2)
                number_content.text = v.split(" ")[0].strip("*")
                price_content = table.cell(j-2,3)
                price_content.text = v.split(" ")[1]
                subtotal_content = table.cell(j-2,4)
                subtotal_content.text = str(int(number_content.text)*int(price_content.text))
                total += int(number_content.text)*int(price_content.text)
            j=j+1
        
        
        for i in range(len(quotation)-2):
            table.cell(i,0).merge(table.cell(i,1))

        for i in range(len(quotation)-2):
            table.cell(i,0).width = Cm(9)
            table.cell(i,1).width = Cm(4.02)
            table.cell(i,2).width = Cm(4.02)
            table.cell(i,3).width = Cm(4.02)
            table.cell(i,4).width = Cm(3.44)
            
        
        for i,row in enumerate(table.rows):
            if i < 2:
                row.height = Cm(0.82)

        table.style = 'Table Grid'
        table_mystyle(table,0)

    table = doc.add_table(rows=12, cols=6)    

    redundent_row = table.cell(11,0)
        
    for i in range(1,6):
        table.cell(0,0).merge(table.cell(0,i))
        table.cell(11,0).merge(table.cell(11,i))
        table.cell(0,0).text = "運費"
        table.cell(0,0).paragraphs[0].runs[0].font.bold = True

    redundent_row.text = """循拾開立電子發票，容器收回及尾款支付後三天內寄送到聯絡人信箱。
    □ 發票需事前提供核銷 (需先支付訂金)  □ 發票分開開立 (請額外告知開立明細)"""

    source = table.cell(1,0)
    source.text = "起始點"
    source_content = table.cell(2,0)
    source_content.text = "循拾公司"
    destination = table.cell(1,1)
    destination.text = "終點"
    destination_content = table.cell(2,1)
    destination_content.text = info["地址"]
    explain = table.cell(1,2)
    explain.text = "說明"
    cost = table.cell(1,3)
    cost.text = "費用"
    cost_content = table.cell(2,3)
    cost_content.text = info["運費"]
    days = table.cell(1,4)
    days.text = "天數"
    days_content = table.cell(2,4)
    days_content.text = str(len(quotation_list[1:]))
    subtotal = table.cell(1,5)
    subtotal.text = "小計"
    subtotal = table.cell(2,5)
    subtotal.text = str(int(info["運費"])*len(quotation_list[1:]))

    for i in range(11):
        table.cell(i,0).width = Cm(3)
        table.cell(i,1).width = Cm(7.46)
        table.cell(i,2).width = Cm(4.02)
        table.cell(i,3).width = Cm(1.93)
        table.cell(i,4).width = Cm(1.51)
        table.cell(i,5).width = Cm(1.52)

    for i,row in enumerate(table.rows):
        if i == 0 or i==3:
            row.height = Cm(0.82)    
        else :
            row.height = Cm(0.69)    
    for i in range(1,6):
        table.cell(3,0).merge(table.cell(3,i))
        table.cell(3,0).text = "加值服務"
        table.cell(3,0).paragraphs[0].runs[0].font.bold = True

    for i in range(4,7):
        for j in range(1,3):
            table.cell(i,0).merge(table.cell(i,j))

    service_item = table.cell(4,0) 
    service_item.text = "服務項目"

    service_item_1 = table.cell(5,0) 
    service_item_1.text = "上下樓"

    service_item_2 = table.cell(6,0) 
    service_item_2.text = "廚餘垃圾代為處理"

    cost = table.cell(4,3)
    cost.text = "費用"
    cost_content_1 = table.cell(5,3)
    cost_content_2 = table.cell(6,3)

    number = table.cell(4,4)
    number.text = "數量"
    number_content_1 = table.cell(5,4)
    number_content_2 = table.cell(6,4)

    subtotal = table.cell(4,5)
    subtotal.text = "小計"
    subtotal_1 = table.cell(5,5)
    subtotal_2 = table.cell(6,5)

    try:
        _ = info["上下樓"] 
        cost_content_1.text = "200"
        number_content_1.text = info["上下樓"]
        total = total + 200*int(info["上下樓"])
        subtotal_1.text = str(200*int(info["上下樓"]))
    except:
        cost_content_1.text = "無"
        number_content_1.text = "無"
        subtotal_1.text ="無"
        info["上下樓"] = 0

    try:
        _ = info["廚餘"] 
        cost_content_2.text = "200"
        number_content_2.text = info["廚餘"]
        total = total + 200*int(info["廚餘"])
        subtotal_2.text = str(200*int(info["廚餘"]))
    except:
        cost_content_2.text = "無"
        number_content_2.text = "無"
        subtotal_2.text ="無"
        info["廚餘"] = 0

    table.style = 'Table Grid'

    total = total + int(info["運費"])*len(quotation_list[1:])

    for i in range(7,11):
        for j in range(1,5):
            table.cell(i,0).merge(table.cell(i,j))

    table.cell(7,0).text = "合計："
    table.cell(8,0).text = "折扣："
    table.cell(9,0).text = "稅額："
    table.cell(10,0).text = "總計："

    table.cell(7,5).text = str(total)
    if "租餐具" in info["類別"]:
        table.cell(9,5).text = str(round(total * 0.05)) 
        total = round(total * 1.05)
        
    else:
        table.cell(9,5).text = str(round(int((int(info["運費"])*len(quotation_list[1:]) +  200*int(info["上下樓"]) + 200*int(info["廚餘"])) * 0.05)))
        total = total+(round(int((int(info["運費"])*len(quotation_list[1:]) +  200*int(info["上下樓"]) + 200*int(info["廚餘"])) * 0.05)))
        

    table.cell(8,5).text = "0"
    table.cell(10,5).text = str(total)
    table_mystyle(table,0)
    for i in range(7,11):
        table.cell(i,0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.cell(11,0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cell_border(
            table.cell(7,0),
            bottom={"sz": 12, "color": "#FFFFFF", "val": "single"},
        )
    set_cell_border(
            table.cell(8,0),
            bottom={"sz": 12, "color": "#FFFFFF", "val": "single"},
        )
    doc.add_paragraph("")
    text = """注意事項:"""
    paragraph = doc.add_paragraph(text)
    paragraph.runs[0].font.bold = True

    for run in paragraph.runs:
        font = run.font
        font.name = '標楷體'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') 

    text = """1. 容器使用完畢後請清空廚餘或飲料殘留，堆疊放回箱/袋內，無加購廚餘代處理服務但殘留廚餘
    過多將酌收處理費用1000元。
    2. 循環容器遺失或是毀損一個150元。
    3. 循拾合作餐廳餐點費用已包含稅額，租借費、服務費等額外計算5% 營業稅。
    4. 活動日期前2天，因私人原因取消訂單將收取全部費用。
    5. 因天災或不可抗因素致使訂單取消，將扣除已購買準備之食材、運送費等相關費用後全數退回，
    若訂金不足以支付將額外收取費用。
    請以電子或列印簽回報價單，代表對內容與金額的確認無誤與了解。"""
    paragraph = doc.add_paragraph(text)
    paragraph.runs[0].font.size = Pt(11)

    for run in paragraph.runs:
        font = run.font
        font.name = '標楷體'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') 


    doc.add_paragraph("")
    text = """費用資訊:"""
    paragraph = doc.add_paragraph(text)
    paragraph.runs[0].font.bold = True

    for run in paragraph.runs:
        font = run.font
        font.name = '標楷體'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') 

    if total >= 5000 and total < 10000:
        text = """1. 匯款訂金5000元後，訂單成立。
    2. 匯款資訊：
    戶名：循拾股份有限公司
    彰化銀行(009)，大安分行(5130)，帳號：51308666680200
    """
    elif total >= 10000:
        text = """1. 匯款50% 訂金"""+str(total//2)+"""元後，訂單成立。
    2. 匯款資訊：
    戶名：循拾股份有限公司
    彰化銀行(009)，大安分行(5130)，帳號：51308666680200
    """
    else:
        text = """1. 匯款"""+str(total)+"""元後，訂單成立。
    2. 匯款資訊：
    戶名：循拾股份有限公司
    彰化銀行(009)，大安分行(5130)，帳號：51308666680200
    """

    paragraph = doc.add_paragraph(text)
    paragraph.runs[0].font.size = Pt(11)

    for run in paragraph.runs:
        font = run.font
        font.name = '標楷體'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') 

    doc.add_paragraph("")
    table = doc.add_table(rows=3, cols=4)

    title = table.cell(0,0)
    title.text = "客戶簽章："
    title_content = table.cell(0,1)

    contact = table.cell(0,2)
    contact.text = "聯絡人："
    contact_content = table.cell(0,3)
    contact_content.text = "陳陳"

    phone = table.cell(1,2)
    phone.text = "Email："
    phone_content = table.cell(1,3)
    phone_content.text = "loopickco@loopick.com.tw"

    mail = table.cell(2,0)
    mail.text = "日期："

    tax = table.cell(2,2)
    tax.text = "電話："
    tax_content = table.cell(2,3)
    tax_content.text = "02-77098711#500"


    table.style = 'Table Grid'
    for i in range(3):
        table.cell(i,1).width = Cm(2.3)

    table_mystyle(table,1)
    table.style = 'Table Normal'
    # Now save the document to a location
    doc.save("output/"+info['報價單號'] +"_"+ info['抬頭'] + ".docx")
    shutil.copyfile('quotation.json', "output/"+info['報價單號'] +"_"+ info['抬頭'] + ".json")

    copy_csv = {
        "顧客名稱(有抬頭後記得替換)":[info["抬頭"]],
        "訂購需求":[""],
        "聯絡管道":[""],
        "統編":[info["統編"]],
        "聯絡窗口":[info["聯絡人"]],
        "連絡電話":[info["電話"]], 
        "聯絡信箱":[info["信箱"]], 
        "總費用": [str(total)]
    }

    df = pd.DataFrame(copy_csv)
    df.to_csv("output/"+info['報價單號'] +"_"+ info['抬頭'] + ".csv", encoding='utf_8_sig')
    filename=info['報價單號'] +"_"+ info['抬頭']
    return filename


if __name__ == "__main__":
    process_quotation(quotation_data)